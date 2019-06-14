#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''-------------------------------------漂亮的分割线----------------------------------------------
作者          ：李艳丽
时间          ：2019/3/20 16:59
项目和文件名称：PyCharm Community Edition WordCompare.py
脚本功能说明  ：两份word的比较
-------------------------------------漂亮的分割线----------------------------------------------'''

import docx
import codecs
from difflib import HtmlDiff
file1=docx.Document('文档1.docx')
file2=docx.Document('文档2.docx')

para1=''
para2=''

for para in file1.paragraphs:
    para1=para1+para.text+'\n'
    print(para1)
for para in file2.paragraphs:
    para2=para2+para.text+'\n'
    print(para2)
delta_html=HtmlDiff().make_file(para1.splitlines(),para2.splitlines())
with codecs.open('diff.html','w',encoding='utf-8') as f:
    f.write(delta_html)
